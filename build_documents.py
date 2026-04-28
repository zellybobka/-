from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
TITLE = "компьютерное структурно-функциональное моделирование волоконно-оптических линий передачи и оценка качества"
DISCIPLINE = "СЕТИ И СИСТЕМЫ ПЕРЕДАЧИ ИНФОРМАЦИИ"


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)


def add_cover_page(document: Document) -> None:
    lines = [
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "Федеральное государственное автономное образовательное учреждение высшего образования",
        "«Санкт-Петербургский государственный университет аэрокосмического приборостроения»",
        "КАФЕДРА 25",
    ]
    for text in lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        if text in {lines[0], lines[3]}:
            run.bold = True

    for _ in range(2):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("ОТЧЕТ \nЗАЩИЩЕН С ОЦЕНКОЙ")

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("ПРЕПОДАВАТЕЛЬ")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ОТЧЕТ ПО ИНДИВИДУАЛЬНОМУ ЗАДАНИЮ №2")
    run.bold = True

    table = document.add_table(rows=3, cols=5)
    table.rows[0].cells[0].text = "профессор, д-р. техн. наук, доцент"
    table.rows[0].cells[4].text = "К. З. Билятдинов"
    table.rows[1].cells[0].text = "должность, уч. степень, звание"
    table.rows[1].cells[2].text = "подпись, дата"
    table.rows[1].cells[4].text = "инициалы, фамилия"

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(TITLE)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"по дисциплине: {DISCIPLINE}")

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("РАБОТУ ВЫПОЛНИЛ")

    table = document.add_table(rows=2, cols=6)
    table.rows[0].cells[0].text = "Студент гр. №"
    table.rows[0].cells[1].text = "3333"
    table.rows[0].cells[5].text = "Зеленин А.И."
    table.rows[1].cells[3].text = "подпись, дата"
    table.rows[1].cells[5].text = "инициалы, фамилия"

    for _ in range(5):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Санкт-Петербург\n2026")
    document.add_page_break()


def add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True


def add_body_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.add_run(text)


def build_abstract() -> Path:
    doc = Document()
    configure_document(doc)
    add_cover_page(doc)
    add_heading(doc, "Реферат")
    paragraphs = [
        "Программа для ЭВМ «Компьютерное структурно-функциональное моделирование ВОЛС и оценка качества» предназначена для сравнения базовой «идеальной» и реальных структурно-функциональных моделей волоконно-оптической линии связи и формирования итогового рейтинга качества вариантов построения ВОЛС.",
        "Приложение реализовано на языке Python в виде веб-программы с использованием библиотек Streamlit, Pandas и NumPy. Доступ к программе осуществляется через браузер по локальному адресу после запуска сервера Streamlit, что обеспечивает использование программы как в операционной системе Windows, так и в Linux.",
        "Функциональные возможности программы включают отображение базовой модели ВОЛС, выбор и сравнение реальных моделей, добавление пользовательской модели, автоматическое формирование матриц BQ, RQ, CQ и весовых коэффициентов ql, проверку соответствия моделей требованиям, расчет комплексного показателя качества эксплуатации Qкэ и построение итогового рейтинга.",
        "В обновленной версии программы используются показатели пропускной способности, затухания сигнала, сквозной задержки, энергопотребления, коэффициента готовности и запаса оптического бюджета, а интерфейс оформлен в стиле Windows XP для повышения наглядности и удобства работы пользователя.",
    ]
    for paragraph in paragraphs:
        add_body_paragraph(doc, paragraph)
    path = BASE_DIR / "Реферат.docx"
    doc.save(path)
    return path


def build_user_guide() -> Path:
    doc = Document()
    configure_document(doc)
    add_cover_page(doc)
    add_heading(doc, "Руководство пользователя")
    paragraphs = [
        "Программа предназначена для оценки качества вариантов ВОЛС путем сравнения базовой модели с выбранными реальными моделями. После запуска приложения в браузере отображаются разделы с базовой моделью, справочными таблицами, блоком добавления пользовательской модели, базовыми матрицами, результатами расчета и итоговым рейтингом.",
        "Для начала работы необходимо установить Python версии 3.10 или выше. В операционной системе Windows следует выполнить файлы install_windows.bat и run_windows.bat. В Linux необходимо запустить install_linux.sh, затем run_linux.sh. После запуска следует открыть в браузере адрес http://localhost:8501.",
        "1. В разделе «Базовая идеальная модель» просмотрите эталонные показатели качества, используемые в расчетах.",
        "2. В разделе справочника ознакомьтесь с перечнем показателей качества и списком моделей ВОЛС.",
        "3. При необходимости в разделе добавления пользовательской модели введите название, описание, допустимый расход ресурсов и отредактируйте значения показателей в таблице, после чего нажмите кнопку «Добавить модель в расчет».",
        "4. В разделе выбора моделей отметьте модели, которые должны участвовать в сравнении с базовой моделью.",
        "5. В разделе базовых матриц просмотрите матрицы BQ, BQ1, BQ2 и ql, сформированные программой автоматически.",
        "6. В разделе результатов расчета откройте нужную модель и просмотрите матрицы RQ, RQ1, RQ2, CQ1, CQ2, CQ и ql·ΔQl. Если модель не соответствует требованиям, программа покажет отрицательные элементы и исключит такую модель из рейтинга.",
        "7. В разделе итогового рейтинга просмотрите упорядоченный список допустимых моделей и при необходимости скачайте текстовый отчет.",
    ]
    for paragraph in paragraphs:
        add_body_paragraph(doc, paragraph)
    path = BASE_DIR / "Руководство_пользователя.docx"
    doc.save(path)
    return path


def build_code_listing() -> Path:
    doc = Document()
    configure_document(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    add_cover_page(doc)
    add_heading(doc, "Листинг исходного кода")

    for file_name in ["vols_quality_app.py", "vols_sfm_app.py"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(file_name)
        run.bold = True

        source = (BASE_DIR / file_name).read_text(encoding="utf-8")
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(source)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)

        doc.add_page_break()

    path = BASE_DIR / "Листинг_кода.docx"
    doc.save(path)
    return path


def build_zip() -> Path:
    archive_path = BASE_DIR / "Комплект_ВОЛС_исправленный.zip"
    include = [
        "README.md",
        "requirements.txt",
        "sample_model.json",
        "vols_quality_app.py",
        "vols_sfm_app.py",
        "install_windows.bat",
        "run_windows.bat",
        "install_linux.sh",
        "run_linux.sh",
        "Реферат.docx",
        "Руководство_пользователя.docx",
        "Листинг_кода.docx",
    ]
    with zipfile.ZipFile(archive_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name in include:
            archive.write(BASE_DIR / name, arcname=name)
    return archive_path


def main() -> None:
    build_abstract()
    build_user_guide()
    build_code_listing()
    build_zip()


if __name__ == "__main__":
    main()
